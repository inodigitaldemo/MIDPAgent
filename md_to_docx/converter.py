"""
md_to_docx.converter
~~~~~~~~~~~~~~~~~~~~
Core conversion helpers that implement the Markdown → HTML → DOCX pipeline.

All functions accept :class:`pathlib.Path` objects (or strings) and return
the :class:`pathlib.Path` of the newly created output file so callers can
chain the steps easily (DRY principle – each step is defined once).
"""

import re
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Default template assets (relative to the repo root)
# ---------------------------------------------------------------------------

_REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_CSS_PATH: Path = _REPO_ROOT / "DocTemplates" / "CSS" / "style.css"
DEFAULT_HTML_TEMPLATE_PATH: Path = _REPO_ROOT / "DocTemplates" / "HTML" / "template.html"

_DEFAULT_BRAND_COLORS = {
    "document-header-bg": "000000",
    "document-footer-bg": "000000",
    "document-header-text": "FFB81C",
    "document-footer-text": "FFB81C",
}


def _normalize_hex_color(value: str | None, fallback: str) -> str:
    """Return a six-character uppercase hex color without a leading ``#``."""
    if not value:
        return fallback

    normalized = value.strip()
    if normalized.startswith("#"):
        normalized = normalized[1:]

    if len(normalized) == 3:
        normalized = "".join(char * 2 for char in normalized)

    if len(normalized) != 6 or any(char not in "0123456789abcdefABCDEF" for char in normalized):
        return fallback

    return normalized.upper()


def _load_css_variables(css_path: Path | str | None) -> dict[str, str]:
    """Extract CSS custom properties from the stylesheet used for HTML output."""
    css_variables = dict(_DEFAULT_BRAND_COLORS)
    if css_path is None:
        return css_variables

    css_path = Path(css_path)
    if not css_path.exists():
        return css_variables

    contents = css_path.read_text(encoding="utf-8")
    for variable_name, raw_value in re.findall(r"--([\w-]+)\s*:\s*([^;]+);", contents):
        if variable_name in css_variables:
            css_variables[variable_name] = _normalize_hex_color(raw_value, css_variables[variable_name])

    return css_variables


def _clear_container(container) -> None:
    """Remove existing paragraphs and tables from a header/footer container."""
    element = container._element
    for child in list(element):
        element.remove(child)


def _set_cell_shading(cell, fill: str) -> None:
    """Apply solid background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()

    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_borders_none(cell) -> None:
    """Remove visible borders from a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()

    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)

    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tc_pr.append(borders)


def _format_cell_paragraph(paragraph, alignment: WD_ALIGN_PARAGRAPH) -> None:
    """Set consistent spacing and alignment for header/footer text."""
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _style_run(run, color_hex: str, size_pt: float, bold: bool = True) -> None:
    """Apply consistent font styling to a run."""
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _append_word_field(paragraph, instruction: str, color_hex: str, size_pt: float) -> None:
    """Append a dynamic Word field (e.g. PAGE) to a paragraph."""
    begin_run = paragraph.add_run()
    _style_run(begin_run, color_hex, size_pt)
    begin_fld = OxmlElement("w:fldChar")
    begin_fld.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin_fld)

    instr_run = paragraph.add_run()
    _style_run(instr_run, color_hex, size_pt)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr_run._r.append(instr_text)

    separate_run = paragraph.add_run()
    _style_run(separate_run, color_hex, size_pt)
    separate_fld = OxmlElement("w:fldChar")
    separate_fld.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate_fld)

    default_value_run = paragraph.add_run("1")
    _style_run(default_value_run, color_hex, size_pt)

    end_run = paragraph.add_run()
    _style_run(end_run, color_hex, size_pt)
    end_fld = OxmlElement("w:fldChar")
    end_fld.set(qn("w:fldCharType"), "end")
    end_run._r.append(end_fld)


def _ensure_section_dimensions(section) -> int:
    """Return usable section width, populating sensible defaults when absent."""
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.left_margin is None:
        section.left_margin = Inches(1)
    if section.right_margin is None:
        section.right_margin = Inches(1)
    if section.top_margin is None:
        section.top_margin = Inches(1)
    if section.bottom_margin is None:
        section.bottom_margin = Inches(1)

    return section.page_width - section.left_margin - section.right_margin


def _add_header_band(section, title_text: str, subtitle_text: str, colors: dict[str, str]) -> None:
    """Create a styled Word header band for a section."""
    header = section.header
    header.is_linked_to_previous = False
    _clear_container(header)

    usable_width = _ensure_section_dimensions(section)
    table = header.add_table(rows=1, cols=1, width=usable_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.rows[0].cells[0]
    cell.width = usable_width
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(cell, colors["document-header-bg"])
    _set_cell_borders_none(cell)

    title_paragraph = cell.paragraphs[0]
    _format_cell_paragraph(title_paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    title_run = title_paragraph.add_run(title_text)
    title_run.font.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor.from_string(colors["document-header-text"])

    subtitle_paragraph = cell.add_paragraph()
    _format_cell_paragraph(subtitle_paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    subtitle_run = subtitle_paragraph.add_run(subtitle_text)
    subtitle_run.font.bold = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor.from_string(colors["document-header-text"])


def _add_footer_band(section, left_text: str, right_text: str, colors: dict[str, str]) -> None:
    """Create a styled Word footer band for a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_container(footer)

    usable_width = _ensure_section_dimensions(section)
    table = footer.add_table(rows=1, cols=2, width=usable_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left_width = int(usable_width * 0.65)
    right_width = usable_width - left_width

    for index, cell in enumerate(table.rows[0].cells):
        cell.width = left_width if index == 0 else right_width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, colors["document-footer-bg"])
        _set_cell_borders_none(cell)

    left_paragraph = table.rows[0].cells[0].paragraphs[0]
    _format_cell_paragraph(left_paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    left_run = left_paragraph.add_run(left_text)
    _style_run(left_run, colors["document-footer-text"], 9.5)

    right_paragraph = table.rows[0].cells[1].paragraphs[0]
    _format_cell_paragraph(right_paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
    right_run = right_paragraph.add_run(right_text)
    _style_run(right_run, colors["document-footer-text"], 9.5)
    separator_run = right_paragraph.add_run(" | Page ")
    _style_run(separator_run, colors["document-footer-text"], 9.5)
    _append_word_field(right_paragraph, " PAGE \\* MERGEFORMAT ", colors["document-footer-text"], 9.5)


def _apply_docx_branding(
    docx_path: Path | str,
    css_path: Path | str | None = DEFAULT_CSS_PATH,
    document_title: str | None = None,
) -> Path:
    """Add real Word header/footer content to the generated DOCX file.

    Pandoc's HTML template affects HTML output, but Word page headers and footers
    must be written into the DOCX section properties (or come from a reference
    document). This post-processing step applies branded header/footer bands to
    every section in the generated DOCX.
    """
    docx_path = Path(docx_path)
    colors = _load_css_variables(css_path)

    document = Document(docx_path)
    resolved_title = document_title or document.core_properties.title or docx_path.stem

    for section in document.sections:
        section.different_first_page_header_footer = False
        section.start_type = WD_SECTION_START.NEW_PAGE if section.start_type is None else section.start_type
        _add_header_band(section, resolved_title, "Controlled Document", colors)
        _add_footer_band(section, "Quality Management System", "Generated from Markdown", colors)

    document.save(docx_path)
    return docx_path


# ---------------------------------------------------------------------------
# Individual conversion steps
# ---------------------------------------------------------------------------

def md_to_html(
    md_path: Path | str,
    html_path: Path | str | None = None,
    css_path: Path | str | None = DEFAULT_CSS_PATH,
    template_path: Path | str | None = DEFAULT_HTML_TEMPLATE_PATH,
) -> Path:
    """Convert a Markdown file to a self-contained HTML file.

    Parameters
    ----------
    md_path:
        Path to the source ``.md`` file.
    html_path:
        Destination path for the generated ``.html`` file.  Defaults to the
        same directory and stem as *md_path* with a ``.html`` extension.
    css_path:
        Path to a CSS stylesheet to embed in the HTML.  Defaults to
        ``DocTemplates/CSS/style.css``.  Pass ``None`` to omit styling.
    template_path:
        Path to the Pandoc HTML template. Defaults to
        ``DocTemplates/HTML/template.html``. Pass ``None`` to use Pandoc's
        built-in template.

    Returns
    -------
    :class:`pathlib.Path`
        The path of the generated HTML file.
    """
    md_path = Path(md_path)

    if html_path is None:
        html_path = md_path.with_suffix(".html")
    html_path = Path(html_path)

    extra_args = ["--standalone"]
    if css_path is not None:
        css_path = Path(css_path)
        if css_path.exists():
            extra_args.extend(["--css", str(css_path)])
    if template_path is not None:
        template_path = Path(template_path)
        if template_path.exists():
            extra_args.extend(["--template", str(template_path)])

    pypandoc.convert_file(
        str(md_path),
        "html",
        outputfile=str(html_path),
        extra_args=extra_args,
    )
    return html_path


def html_to_docx(
    html_path: Path | str,
    docx_path: Path | str | None = None,
    css_path: Path | str | None = DEFAULT_CSS_PATH,
) -> Path:
    """Convert an HTML file to a DOCX file.

    Parameters
    ----------
    html_path:
        Path to the source ``.html`` file.
    docx_path:
        Destination path for the generated ``.docx`` file.  Defaults to the
        same directory and stem as *html_path* with a ``.docx`` extension.
    css_path:
        CSS stylesheet containing brand color variables used to build the Word
        header and footer bands.

    Returns
    -------
    :class:`pathlib.Path`
        The path of the generated DOCX file.
    """
    html_path = Path(html_path)

    if docx_path is None:
        docx_path = html_path.with_suffix(".docx")
    docx_path = Path(docx_path)

    pypandoc.convert_file(str(html_path), "docx", outputfile=str(docx_path))
    _apply_docx_branding(docx_path, css_path=css_path, document_title=html_path.stem)
    return docx_path


# ---------------------------------------------------------------------------
# Convenience wrapper
# ---------------------------------------------------------------------------

def convert_md_to_docx(
    md_path: Path | str,
    output_dir: Path | str | None = None,
    css_path: Path | str | None = DEFAULT_CSS_PATH,
    template_path: Path | str | None = DEFAULT_HTML_TEMPLATE_PATH,
) -> tuple[Path, Path]:
    """Convert a Markdown file to DOCX via an intermediate HTML file.

    The two-step pipeline is:
    1. ``.md``  →  ``.html``  (via :func:`md_to_html`)
    2. ``.html`` → ``.docx``  (via :func:`html_to_docx`)

    Parameters
    ----------
    md_path:
        Path to the source ``.md`` file.
    output_dir:
        Directory in which to write the ``.html`` and ``.docx`` files.
        Defaults to the same directory as *md_path*.
    css_path:
        CSS stylesheet for the HTML step.  See :func:`md_to_html`.
    template_path:
        HTML template for the HTML step. See :func:`md_to_html`.

    Returns
    -------
    tuple[Path, Path]
        ``(html_path, docx_path)`` – the paths of the generated files.
    """
    md_path = Path(md_path)

    if output_dir is None:
        output_dir = md_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    html_path = output_dir / md_path.with_suffix(".html").name
    docx_path = output_dir / md_path.with_suffix(".docx").name

    md_to_html(md_path, html_path, css_path, template_path)
    html_to_docx(html_path, docx_path, css_path)

    return html_path, docx_path
